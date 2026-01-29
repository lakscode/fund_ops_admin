"""
Generate comprehensive documentation for Fund Ops Admin application
Creates a DOCX file with detailed information about UI and API modules
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def add_heading(doc, text, level=1):
    """Add a formatted heading"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_paragraph(doc, text, bold=False):
    """Add a formatted paragraph"""
    p = doc.add_paragraph(text)
    if bold:
        p.runs[0].bold = True
    return p

def add_table_row(table, cells_data, bold_first=False):
    """Add a row to table"""
    row = table.add_row()
    for i, cell_data in enumerate(cells_data):
        row.cells[i].text = str(cell_data)
        if bold_first and i == 0:
            row.cells[i].paragraphs[0].runs[0].bold = True

def create_documentation():
    """Generate the complete documentation"""
    doc = Document()

    # Title
    title = doc.add_heading('Fund Operations Admin', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('System Documentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)

    date_para = doc.add_paragraph(f'Generated on: {datetime.now().strftime("%B %d, %Y")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ============================================================================
    # TABLE OF CONTENTS
    # ============================================================================
    add_heading(doc, 'Table of Contents', 1)
    doc.add_paragraph('1. Overview')
    doc.add_paragraph('2. Technology Stack')
    doc.add_paragraph('3. Frontend (UI) Documentation')
    doc.add_paragraph('   3.1 Dependencies')
    doc.add_paragraph('   3.2 Module Structure')
    doc.add_paragraph('4. Backend (API) Documentation')
    doc.add_paragraph('   4.1 Dependencies')
    doc.add_paragraph('   4.2 Module Structure')
    doc.add_paragraph('5. Database Configuration')
    doc.add_paragraph('6. Authentication & Security')
    doc.add_paragraph('7. Import/Export Functionality')
    doc.add_paragraph('8. API Endpoints')

    doc.add_page_break()

    # ============================================================================
    # OVERVIEW
    # ============================================================================
    add_heading(doc, '1. Overview', 1)
    add_paragraph(doc, 'Fund Operations Admin is a comprehensive web application designed to manage '
                       'investment funds, investors, properties, and organizations. The system provides '
                       'a modern interface for fund managers, administrators, and other stakeholders to '
                       'track and manage financial operations.')

    doc.add_paragraph()
    add_paragraph(doc, 'Key Features:', bold=True)
    doc.add_paragraph('• Multi-organization support with role-based access control')
    doc.add_paragraph('• Fund management and tracking')
    doc.add_paragraph('• Investor portfolio management')
    doc.add_paragraph('• Property asset tracking')
    doc.add_paragraph('• User and organization administration')
    doc.add_paragraph('• Secure authentication with JWT tokens')

    doc.add_page_break()

    # ============================================================================
    # TECHNOLOGY STACK
    # ============================================================================
    add_heading(doc, '2. Technology Stack', 1)

    add_heading(doc, 'Frontend Technologies', 2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Technology'
    header_cells[1].text = 'Version/Description'

    frontend_tech = [
        ('React', '18.2.0 - UI Framework'),
        ('TypeScript', '5.3.3 - Type-safe JavaScript'),
        ('Vite', '5.0.10 - Build tool and dev server'),
        ('React Router', '6.21.0 - Client-side routing'),
        ('Tailwind CSS', '3.4.0 - Utility-first CSS framework'),
        ('Axios', '1.6.2 - HTTP client'),
        ('Lucide React', '0.294.0 - Icon library')
    ]

    for tech, desc in frontend_tech:
        add_table_row(table, [tech, desc], bold_first=True)

    doc.add_paragraph()

    add_heading(doc, 'Backend Technologies', 2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Technology'
    header_cells[1].text = 'Version/Description'

    backend_tech = [
        ('FastAPI', '0.109.0 - Modern Python web framework'),
        ('Uvicorn', '0.27.0 - ASGI server'),
        ('MongoDB', 'Atlas - NoSQL database'),
        ('Motor', '3.3.2 - Async MongoDB driver'),
        ('Pydantic', '2.5.3 - Data validation'),
        ('Python-Jose', '3.3.0 - JWT token handling'),
        ('Passlib', '1.7.4 - Password hashing')
    ]

    for tech, desc in backend_tech:
        add_table_row(table, [tech, desc], bold_first=True)

    doc.add_page_break()

    # ============================================================================
    # FRONTEND DOCUMENTATION
    # ============================================================================
    add_heading(doc, '3. Frontend (UI) Documentation', 1)

    add_heading(doc, '3.1 Dependencies', 2)

    add_heading(doc, 'Core Dependencies', 3)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Package'
    header_cells[1].text = 'Version'
    header_cells[2].text = 'Purpose'

    core_deps = [
        ('react', '^18.2.0', 'Core React library for building UI components'),
        ('react-dom', '^18.2.0', 'React rendering for web browsers'),
        ('react-router-dom', '^6.21.0', 'Declarative routing for React applications'),
        ('axios', '^1.6.2', 'Promise-based HTTP client for API requests'),
        ('lucide-react', '^0.294.0', 'Beautiful, consistent icon library'),
        ('typescript', '^5.3.3', 'Static type checking for JavaScript')
    ]

    for pkg, ver, purpose in core_deps:
        add_table_row(table, [pkg, ver, purpose])

    doc.add_paragraph()

    add_heading(doc, 'Development Dependencies', 3)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Package'
    header_cells[1].text = 'Version'
    header_cells[2].text = 'Purpose'

    dev_deps = [
        ('@vitejs/plugin-react', '^4.2.1', 'Official Vite plugin for React with Fast Refresh'),
        ('vite', '^5.0.10', 'Next-generation frontend build tool'),
        ('autoprefixer', '^10.4.16', 'PostCSS plugin to add vendor prefixes'),
        ('postcss', '^8.4.32', 'Tool for transforming CSS with JavaScript'),
        ('tailwindcss', '^3.4.0', 'Utility-first CSS framework')
    ]

    for pkg, ver, purpose in dev_deps:
        add_table_row(table, [pkg, ver, purpose])

    doc.add_page_break()

    # ============================================================================
    # FRONTEND MODULE STRUCTURE
    # ============================================================================
    add_heading(doc, '3.2 Module Structure', 2)

    add_heading(doc, 'Pages Module (src/pages/)', 3)
    add_paragraph(doc, 'Contains main application pages and views.')
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'File'
    header_cells[1].text = 'Description'

    pages = [
        ('Login.tsx', 'User authentication page with login form and JWT token handling'),
        ('Dashboard.tsx', 'Main dashboard displaying key metrics (funds, investors, properties, allocations)'),
        ('Funds.tsx', 'Fund management page with CRUD operations for investment funds'),
        ('Investors.tsx', 'Investor management interface with portfolio tracking'),
        ('Properties.tsx', 'Property asset management and tracking page')
    ]

    for file, desc in pages:
        add_table_row(table, [file, desc], bold_first=True)

    doc.add_paragraph()

    add_heading(doc, 'Components Module (src/components/)', 3)
    add_paragraph(doc, 'Reusable UI components used across the application.')
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Component'
    header_cells[1].text = 'Description'

    components = [
        ('Layout.tsx', 'Main layout wrapper providing consistent page structure with header and sidebar'),
        ('Header.tsx', 'Top navigation bar with user info and organization selector'),
        ('Sidebar.tsx', 'Side navigation menu with links to all main pages'),
        ('Modal.tsx', 'Reusable modal dialog component for forms and confirmations'),
        ('OrganizationDropdown.tsx', 'Dropdown to switch between organizations user has access to')
    ]

    for comp, desc in components:
        add_table_row(table, [comp, desc], bold_first=True)

    doc.add_paragraph()

    add_heading(doc, 'Context Providers (src/contexts/)', 3)
    add_paragraph(doc, 'React Context providers for global state management.')
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Context'
    header_cells[1].text = 'Description'

    contexts = [
        ('AuthContext.tsx', 'Manages user authentication state, login/logout, and protected routes'),
        ('OrganizationContext.tsx', 'Manages current organization context and user role within organization')
    ]

    for ctx, desc in contexts:
        add_table_row(table, [ctx, desc], bold_first=True)

    doc.add_paragraph()

    add_heading(doc, 'Services Module (src/services/)', 3)
    add_paragraph(doc, 'API service layer for backend communication.')
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Service'
    header_cells[1].text = 'Description'

    services = [
        ('api.ts', 'Base Axios configuration with interceptors for authentication headers'),
        ('auth.ts', 'Authentication services (login, logout, token management, user info)'),
        ('funds.ts', 'Fund CRUD operations and organization-specific fund queries'),
        ('investors.ts', 'Investor management and portfolio allocation operations'),
        ('properties.ts', 'Property asset CRUD operations'),
        ('organizations.ts', 'Organization management and user-organization associations')
    ]

    for svc, desc in services:
        add_table_row(table, [svc, desc], bold_first=True)

    doc.add_paragraph()

    add_heading(doc, 'Types Module (src/types/)', 3)
    add_paragraph(doc, 'TypeScript type definitions and interfaces.')
    doc.add_paragraph()
    add_paragraph(doc, 'index.ts - Centralized type definitions for:', bold=True)
    doc.add_paragraph('• User, Organization, UserOrganization')
    doc.add_paragraph('• Fund, Investor, InvestorFund')
    doc.add_paragraph('• Property, AuthResponse, LoginRequest')

    doc.add_page_break()

    # ============================================================================
    # BACKEND DOCUMENTATION
    # ============================================================================
    add_heading(doc, '4. Backend (API) Documentation', 1)

    add_heading(doc, '4.1 Dependencies', 2)

    add_heading(doc, 'Core Dependencies', 3)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Package'
    header_cells[1].text = 'Version'
    header_cells[2].text = 'Purpose'

    api_core_deps = [
        ('fastapi', '0.109.0', 'Modern, fast web framework for building APIs with Python'),
        ('uvicorn', '0.27.0', 'Lightning-fast ASGI server implementation'),
        ('pydantic', '2.5.3', 'Data validation using Python type annotations'),
        ('pydantic-settings', '2.1.0', 'Settings management using Pydantic'),
        ('python-dotenv', '1.0.0', 'Read environment variables from .env file')
    ]

    for pkg, ver, purpose in api_core_deps:
        add_table_row(table, [pkg, ver, purpose])

    doc.add_paragraph()

    add_heading(doc, 'Database Dependencies', 3)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Package'
    header_cells[1].text = 'Version'
    header_cells[2].text = 'Purpose'

    db_deps = [
        ('motor', '3.3.2', 'Asynchronous Python driver for MongoDB'),
        ('pymongo', '4.6.1', 'Official MongoDB driver for Python')
    ]

    for pkg, ver, purpose in db_deps:
        add_table_row(table, [pkg, ver, purpose])

    doc.add_paragraph()

    add_heading(doc, 'Authentication Dependencies', 3)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Package'
    header_cells[1].text = 'Version'
    header_cells[2].text = 'Purpose'

    auth_deps = [
        ('passlib[bcrypt]', '1.7.4', 'Password hashing library with bcrypt support'),
        ('python-jose[cryptography]', '3.3.0', 'JWT token creation and validation'),
        ('bcrypt', '4.0.1', 'Strong password hashing algorithm'),
        ('python-multipart', '0.0.6', 'Support for multipart/form-data requests')
    ]

    for pkg, ver, purpose in auth_deps:
        add_table_row(table, [pkg, ver, purpose])

    doc.add_paragraph()

    add_heading(doc, 'Import/Export Dependencies', 3)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Package'
    header_cells[1].text = 'Version'
    header_cells[2].text = 'Purpose'

    export_deps = [
        ('openpyxl', '3.1.2', 'Excel file import/export functionality (XLSX format)')
    ]

    for pkg, ver, purpose in export_deps:
        add_table_row(table, [pkg, ver, purpose])

    doc.add_page_break()

    # ============================================================================
    # BACKEND MODULE STRUCTURE
    # ============================================================================
    add_heading(doc, '4.2 Module Structure', 2)

    add_heading(doc, 'Main Application (main.py)', 3)
    add_paragraph(doc, 'Entry point for the FastAPI application. Configures CORS, includes all routers, '
                       'and defines health check endpoints.')
    doc.add_paragraph()

    add_heading(doc, 'Routes Module (routes/)', 3)
    add_paragraph(doc, 'API endpoint definitions organized by resource.')
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Route File'
    header_cells[1].text = 'Endpoints'

    routes = [
        ('auth.py', 'POST /login - User authentication\nGET /me - Get current user info'),
        ('users.py', 'GET /users/ - List all users\nPOST /users/ - Create user\n'
                     'GET /users/{id} - Get user\nPUT /users/{id} - Update user\n'
                     'DELETE /users/{id} - Delete user\nGET /users/{id}/organizations - User orgs'),
        ('organizations.py', 'GET /organizations/ - List organizations\nPOST /organizations/ - Create org\n'
                            'GET /organizations/{id} - Get org\nPUT /organizations/{id} - Update org\n'
                            'DELETE /organizations/{id} - Delete org'),
        ('funds.py', 'GET /funds/ - List funds\nGET /funds/organization/{org_id} - Org funds\n'
                    'POST /funds/ - Create fund\nGET /funds/{id} - Get fund\n'
                    'PUT /funds/{id} - Update fund\nDELETE /funds/{id} - Delete fund'),
        ('investors.py', 'GET /investors/ - List investors\nGET /investors/organization/{org_id} - Org investors\n'
                        'POST /investors/ - Create investor\nGET /investors/{id} - Get investor\n'
                        'PUT /investors/{id} - Update investor\nDELETE /investors/{id} - Delete investor'),
        ('properties.py', 'GET /properties/ - List properties\nGET /properties/organization/{org_id} - Org properties\n'
                         'POST /properties/ - Create property\nGET /properties/{id} - Get property\n'
                         'PUT /properties/{id} - Update property\nDELETE /properties/{id} - Delete property'),
        ('investor_funds.py', 'GET /investor-funds/ - List allocations\n'
                             'GET /investor-funds/organization/{org_id} - Org allocations\n'
                             'POST /investor-funds/ - Create allocation\n'
                             'GET /investor-funds/{id} - Get allocation'),
        ('user_organizations.py', 'GET /user-organizations/ - List user-org mappings\n'
                                 'POST /user-organizations/ - Create mapping\n'
                                 'GET /user-organizations/{id} - Get mapping\n'
                                 'DELETE /user-organizations/{id} - Delete mapping')
    ]

    for route, endpoints in routes:
        add_table_row(table, [route, endpoints], bold_first=True)

    doc.add_paragraph()

    add_heading(doc, 'Models Module (models/)', 3)
    add_paragraph(doc, 'Pydantic models for data validation and serialization.')
    doc.add_paragraph()
    add_paragraph(doc, 'models.py - Core data models:', bold=True)
    doc.add_paragraph('• Organization, Fund, Investor, Property')
    doc.add_paragraph('• User, UserOrganization, InvestorFund')

    doc.add_paragraph()

    add_heading(doc, 'Schemas Module (schemas/)', 3)
    add_paragraph(doc, 'Request and response schemas for API endpoints.')
    doc.add_paragraph()
    add_paragraph(doc, 'schemas.py - Contains:', bold=True)
    doc.add_paragraph('• Create/Update schemas for all entities')
    doc.add_paragraph('• Response schemas with computed fields')
    doc.add_paragraph('• Login request/response schemas')

    doc.add_paragraph()

    add_heading(doc, 'Database Module (database/)', 3)
    add_paragraph(doc, 'MongoDB connection and initialization.')
    doc.add_paragraph()
    add_paragraph(doc, 'database.py - Functions:', bold=True)
    doc.add_paragraph('• init_database() - Initialize MongoDB connection')
    doc.add_paragraph('• get_db() - Get database instance')
    doc.add_paragraph('• close_connection() - Close database connection')

    doc.add_paragraph()

    add_heading(doc, 'Repositories Module (repositories/)', 3)
    add_paragraph(doc, 'Data access layer for MongoDB operations.')
    doc.add_paragraph()
    add_paragraph(doc, 'mongo_repository.py - Generic repository pattern:', bold=True)
    doc.add_paragraph('• get_all() - Retrieve all documents with pagination')
    doc.add_paragraph('• get_by_id() - Retrieve document by ID')
    doc.add_paragraph('• create() - Insert new document')
    doc.add_paragraph('• update() - Update existing document')
    doc.add_paragraph('• delete() - Delete document')
    doc.add_paragraph('• get_by_field() - Query by specific field')

    doc.add_paragraph()

    add_heading(doc, 'Authentication Module (auth/)', 3)
    add_paragraph(doc, 'Security and authentication utilities.')
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'File'
    header_cells[1].text = 'Description'

    auth_modules = [
        ('utils.py', 'Password hashing, token creation/validation, user authentication'),
        ('permissions.py', 'Role-based access control and permission checks')
    ]

    for file, desc in auth_modules:
        add_table_row(table, [file, desc], bold_first=True)

    doc.add_paragraph()

    add_heading(doc, 'Utilities Module (utils/)', 3)
    add_paragraph(doc, 'Utility functions for common operations.')
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'File'
    header_cells[1].text = 'Description'

    utils_modules = [
        ('import_export.py', 'Import and export functionality for XLSX and JSON formats\n'
                            'Functions: export_to_xlsx(), export_to_json(), import_from_xlsx(), import_from_json()\n'
                            'Also includes data validation, template generation, and field type conversion')
    ]

    for file, desc in utils_modules:
        add_table_row(table, [file, desc], bold_first=True)

    doc.add_paragraph()

    add_heading(doc, 'Configuration (config.py)', 3)
    add_paragraph(doc, 'Application settings and environment configuration.')
    doc.add_paragraph()
    add_paragraph(doc, 'Settings class with:', bold=True)
    doc.add_paragraph('• MongoDB connection parameters')
    doc.add_paragraph('• JWT secret key and algorithm')
    doc.add_paragraph('• Token expiration time')
    doc.add_paragraph('• Environment-based configuration')

    doc.add_paragraph()

    add_heading(doc, 'Data Seeding (seed_data.py)', 3)
    add_paragraph(doc, 'Database initialization script that creates sample data.')
    doc.add_paragraph()
    add_paragraph(doc, 'Generates:', bold=True)
    doc.add_paragraph('• 8 organizations')
    doc.add_paragraph('• 34 users (30 regular + 4 admin users)')
    doc.add_paragraph('• 32 funds across organizations')
    doc.add_paragraph('• 50 investors')
    doc.add_paragraph('• 107 investor-fund allocations')
    doc.add_paragraph('• 40 properties')
    doc.add_paragraph()
    add_paragraph(doc, 'Admin Users Created:', bold=True)
    doc.add_paragraph('• superadmin - Platform super admin (is_superuser=true)')
    doc.add_paragraph('• sysadmin - System admin with admin role in ALL organizations')
    doc.add_paragraph('• orgadmin - Organization admin with admin role in first 2 organizations')
    doc.add_paragraph('• viewer - Regular user with viewer role (read-only access)')

    doc.add_page_break()

    # ============================================================================
    # DATABASE CONFIGURATION
    # ============================================================================
    add_heading(doc, '5. Database Configuration', 1)

    add_heading(doc, 'MongoDB Atlas', 2)
    add_paragraph(doc, 'The application uses MongoDB Atlas as the cloud database solution.')
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Parameter'
    header_cells[1].text = 'Value'

    db_config = [
        ('Cluster', 'cluster0.psvlck1.mongodb.net'),
        ('Database Name', 'Fund_Ops_Admin'),
        ('Connection Type', 'MongoDB Atlas (SRV)'),
        ('Driver', 'Motor (Async) / PyMongo (Sync)'),
        ('Authentication', 'Username/Password with encrypted connection')
    ]

    for param, value in db_config:
        add_table_row(table, [param, value], bold_first=True)

    doc.add_paragraph()

    add_heading(doc, 'Collections', 2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Collection'
    header_cells[1].text = 'Description'

    collections = [
        ('organizations', 'Investment management firms and companies'),
        ('users', 'System users with authentication credentials'),
        ('user_organizations', 'Many-to-many mapping between users and organizations with roles'),
        ('funds', 'Investment funds managed by organizations'),
        ('investors', 'Individual or institutional investors'),
        ('investor_funds', 'Investment allocations between investors and funds'),
        ('properties', 'Real estate and property assets')
    ]

    for coll, desc in collections:
        add_table_row(table, [coll, desc], bold_first=True)

    doc.add_paragraph()

    add_heading(doc, 'Indexes', 2)
    add_paragraph(doc, 'Optimized database indexes for query performance:')
    doc.add_paragraph('• users.username (unique)')
    doc.add_paragraph('• users.email (unique)')
    doc.add_paragraph('• organizations.code (unique)')
    doc.add_paragraph('• user_organizations (user_id, organization_id) - compound unique')
    doc.add_paragraph('• funds.organization_id')
    doc.add_paragraph('• investors.organization_id')
    doc.add_paragraph('• investor_funds (investor_id, fund_id) - compound unique')
    doc.add_paragraph('• properties.fund_id')

    doc.add_page_break()

    # ============================================================================
    # AUTHENTICATION & SECURITY
    # ============================================================================
    add_heading(doc, '6. Authentication & Security', 1)

    add_heading(doc, 'JWT Authentication', 2)
    add_paragraph(doc, 'The application uses JSON Web Tokens (JWT) for secure authentication.')
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Feature'
    header_cells[1].text = 'Implementation'

    auth_features = [
        ('Token Type', 'Bearer JWT tokens'),
        ('Algorithm', 'HS256 (HMAC with SHA-256)'),
        ('Token Expiration', '30 minutes'),
        ('Storage', 'localStorage in browser'),
        ('Password Hashing', 'bcrypt with salt')
    ]

    for feature, impl in auth_features:
        add_table_row(table, [feature, impl], bold_first=True)

    doc.add_paragraph()

    add_heading(doc, 'User Roles', 2)
    add_paragraph(doc, 'The system supports two levels of access control:')
    doc.add_paragraph()
    add_paragraph(doc, 'Platform Level:', bold=True)
    doc.add_paragraph('• is_superuser - Platform super admin with access to all organizations and data')
    doc.add_paragraph()
    add_paragraph(doc, 'Organization-Level Roles:', bold=True)
    doc.add_paragraph('• admin - Organization admin, can manage users within their organization')
    doc.add_paragraph('• fund_manager - Can create and manage funds')
    doc.add_paragraph('• analyst - Can view and analyze data, no modification rights')
    doc.add_paragraph('• viewer - Read-only access to organization data')

    doc.add_paragraph()

    add_heading(doc, 'Security Features', 2)
    doc.add_paragraph('• CORS enabled for cross-origin requests')
    doc.add_paragraph('• Password encryption using bcrypt')
    doc.add_paragraph('• JWT token validation on protected routes')
    doc.add_paragraph('• Automatic token refresh mechanism')
    doc.add_paragraph('• Protected API endpoints requiring authentication')
    doc.add_paragraph('• Organization-scoped data access')

    doc.add_paragraph()

    add_heading(doc, 'Test Credentials', 2)
    add_paragraph(doc, 'All admin accounts use password: admin123')
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Username'
    header_cells[1].text = 'Account Type'
    header_cells[2].text = 'Access Level'

    test_creds = [
        ('superadmin', 'Platform Super Admin', 'Full platform access (is_superuser=true)'),
        ('sysadmin', 'System Admin', 'Admin role in ALL organizations'),
        ('orgadmin', 'Organization Admin', 'Admin role in first 2 organizations'),
        ('viewer', 'Regular User', 'Viewer role (read-only) in first organization'),
        ('[firstname.lastname]', 'Regular Users', 'Various roles, password: password123')
    ]

    for username, acc_type, access in test_creds:
        add_table_row(table, [username, acc_type, access])

    doc.add_page_break()

    # ============================================================================
    # IMPORT/EXPORT FUNCTIONALITY
    # ============================================================================
    add_heading(doc, '7. Import/Export Functionality', 1)

    add_paragraph(doc, 'The application supports bulk import and export of data in XLSX and JSON formats.')
    doc.add_paragraph()

    add_heading(doc, 'Supported Formats', 2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Format'
    header_cells[1].text = 'Description'

    formats = [
        ('XLSX (Excel)', 'Excel workbook format with headers in first row'),
        ('JSON', 'JavaScript Object Notation with array of objects')
    ]

    for fmt, desc in formats:
        add_table_row(table, [fmt, desc], bold_first=True)

    doc.add_paragraph()

    add_heading(doc, 'Export Features', 2)
    doc.add_paragraph('• Organization-scoped exports (users can only export their org data)')
    doc.add_paragraph('• Format selection (XLSX or JSON)')
    doc.add_paragraph('• Automatic field filtering and formatting')
    doc.add_paragraph('• Datetime fields converted to ISO format')
    doc.add_paragraph('• Auto-adjusted column widths for XLSX')

    doc.add_paragraph()

    add_heading(doc, 'Import Features', 2)
    doc.add_paragraph('• File upload with multipart/form-data')
    doc.add_paragraph('• Automatic format detection')
    doc.add_paragraph('• Field validation and type conversion')
    doc.add_paragraph('• Detailed error reporting per row')
    doc.add_paragraph('• Rollback on validation errors')
    doc.add_paragraph('• Template generation for correct format')

    doc.add_paragraph()

    add_heading(doc, 'Supported Entities', 2)
    doc.add_paragraph('The following entities support import/export:')
    doc.add_paragraph('• Funds')
    doc.add_paragraph('• Investors')
    doc.add_paragraph('• Properties')

    doc.add_page_break()

    # ============================================================================
    # API ENDPOINTS
    # ============================================================================
    add_heading(doc, '8. API Endpoints', 1)

    add_paragraph(doc, 'Base URL: http://localhost:8000/api/v1')
    doc.add_paragraph()

    add_heading(doc, 'Authentication Endpoints', 2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Method'
    header_cells[1].text = 'Endpoint'
    header_cells[2].text = 'Description'

    auth_endpoints = [
        ('POST', '/auth/login', 'User login - returns JWT token'),
        ('GET', '/auth/me', 'Get current user info (requires auth)')
    ]

    for method, endpoint, desc in auth_endpoints:
        add_table_row(table, [method, endpoint, desc])

    doc.add_paragraph()

    add_heading(doc, 'Organization Endpoints', 2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Method'
    header_cells[1].text = 'Endpoint'
    header_cells[2].text = 'Description'

    org_endpoints = [
        ('GET', '/organizations/', 'List all organizations'),
        ('POST', '/organizations/', 'Create new organization'),
        ('GET', '/organizations/{id}', 'Get organization by ID'),
        ('PUT', '/organizations/{id}', 'Update organization'),
        ('DELETE', '/organizations/{id}', 'Delete organization')
    ]

    for method, endpoint, desc in org_endpoints:
        add_table_row(table, [method, endpoint, desc])

    doc.add_paragraph()

    add_heading(doc, 'Fund Endpoints', 2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Method'
    header_cells[1].text = 'Endpoint'
    header_cells[2].text = 'Description'

    fund_endpoints = [
        ('GET', '/funds/', 'List all funds'),
        ('GET', '/funds/organization/{org_id}', 'Get funds by organization'),
        ('POST', '/funds/', 'Create new fund'),
        ('GET', '/funds/{id}', 'Get fund by ID'),
        ('PUT', '/funds/{id}', 'Update fund'),
        ('DELETE', '/funds/{id}', 'Delete fund'),
        ('GET', '/funds/export?organization_id={id}', 'Export funds to XLSX/JSON'),
        ('POST', '/funds/import?organization_id={id}', 'Import funds from XLSX/JSON file')
    ]

    for method, endpoint, desc in fund_endpoints:
        add_table_row(table, [method, endpoint, desc])

    doc.add_paragraph()

    add_heading(doc, 'Investor Endpoints', 2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Method'
    header_cells[1].text = 'Endpoint'
    header_cells[2].text = 'Description'

    investor_endpoints = [
        ('GET', '/investors/', 'List all investors'),
        ('GET', '/investors/organization/{org_id}', 'Get investors by organization'),
        ('POST', '/investors/', 'Create new investor'),
        ('GET', '/investors/{id}', 'Get investor by ID'),
        ('PUT', '/investors/{id}', 'Update investor'),
        ('DELETE', '/investors/{id}', 'Delete investor'),
        ('GET', '/investors/export?organization_id={id}', 'Export investors to XLSX/JSON'),
        ('POST', '/investors/import?organization_id={id}', 'Import investors from XLSX/JSON file')
    ]

    for method, endpoint, desc in investor_endpoints:
        add_table_row(table, [method, endpoint, desc])

    doc.add_paragraph()

    add_heading(doc, 'Property Endpoints', 2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Method'
    header_cells[1].text = 'Endpoint'
    header_cells[2].text = 'Description'

    property_endpoints = [
        ('GET', '/properties/', 'List all properties'),
        ('GET', '/properties/organization/{org_id}', 'Get properties by organization'),
        ('POST', '/properties/', 'Create new property'),
        ('GET', '/properties/{id}', 'Get property by ID'),
        ('PUT', '/properties/{id}', 'Update property'),
        ('DELETE', '/properties/{id}', 'Delete property'),
        ('GET', '/properties/export?organization_id={id}', 'Export properties to XLSX/JSON'),
        ('POST', '/properties/import?organization_id={id}', 'Import properties from XLSX/JSON file')
    ]

    for method, endpoint, desc in property_endpoints:
        add_table_row(table, [method, endpoint, desc])

    doc.add_paragraph()

    add_heading(doc, 'Investor-Fund Allocation Endpoints', 2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Method'
    header_cells[1].text = 'Endpoint'
    header_cells[2].text = 'Description'

    allocation_endpoints = [
        ('GET', '/investor-funds/', 'List all allocations'),
        ('GET', '/investor-funds/organization/{org_id}', 'Get allocations by organization'),
        ('POST', '/investor-funds/', 'Create new allocation'),
        ('GET', '/investor-funds/{id}', 'Get allocation by ID')
    ]

    for method, endpoint, desc in allocation_endpoints:
        add_table_row(table, [method, endpoint, desc])

    doc.add_paragraph()
    doc.add_paragraph()

    # Footer
    doc.add_paragraph('_' * 80)
    footer = doc.add_paragraph('End of Documentation')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save document
    filename = 'Fund_Ops_Admin_Documentation_Updated.docx'
    doc.save(filename)
    print(f"Documentation generated successfully: {filename}")

if __name__ == "__main__":
    create_documentation()
